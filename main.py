!apt update && apt install -y ffmpeg

# Yüklendikten sonra dosya adını buraya girin
input_file = "ses_dosyasi.mp4" # Kendi dosya adınızla değiştirin

# Toplam süre (saniye cinsinden)
total_duration_seconds = (54 * 60) + 25


# İlk 3 parçayı 15 dakika, son parçayı kalan süre olarak ayarlayalım.

# Parça süreleri (saniye cinsinden)
duration_part1 = 15 * 60
duration_part2 = 15 * 60
duration_part3 = 15 * 60

# Son parça kalan süre olacak
duration_part4 = total_duration_seconds - (duration_part1 + duration_part2 + duration_part3)

# Başlangıç zamanları (saniye cinsinden)
start_time_part1 = 0
start_time_part2 = duration_part1
start_time_part3 = duration_part1 + duration_part2
start_time_part4 = duration_part1 + duration_part2 + duration_part3


print(f"Bölüm 1 süresi: {duration_part1 // 60} dakika {duration_part1 % 60} saniye")
print(f"Bölüm 2 süresi: {duration_part2 // 60} dakika {duration_part2 % 60} saniye")
print(f"Bölüm 3 süresi: {duration_part3 // 60} dakika {duration_part3 % 60} saniye")
print(f"Bölüm 4 süresi: {duration_part4 // 60} dakika {duration_part4 % 60} saniye")


# FFmpeg komutları
# -ss: başlangıç zamanı
# -t: süre
# -i: giriş dosyası
# -acodec pcm_s16le: WAV için kodek
# -ar 44100: örnekleme oranı (genel standart)
# -ac 2: kanal sayısı (stereo)

# Bölüm 1
!ffmpeg -ss {start_time_part1} -t {duration_part1} -i "{input_file}" -acodec pcm_s16le -ar 44100 -ac 2 output_part1.wav

# Bölüm 2
!ffmpeg -ss {start_time_part2} -t {duration_part2} -i "{input_file}" -acodec pcm_s16le -ar 44100 -ac 2 output_part2.wav

# Bölüm 3
!ffmpeg -ss {start_time_part3} -t {duration_part3} -i "{input_file}" -acodec pcm_s16le -ar 44100 -ac 2 output_part3.wav

# Bölüm 4
!ffmpeg -ss {start_time_part4} -t {duration_part4} -i "{input_file}" -acodec pcm_s16le -ar 44100 -ac 2 output_part4.wav

print("\nBölme ve dönüştürme işlemi tamamlandı!")
print("output_part1.wav, output_part2.wav, output_part3.wav ve output_part4.wav dosyaları oluşturuldu.")

# Oluşturulan dosyaları indirme (isteğe bağlı)
from google.colab import files
files.download('output_part1.wav')
files.download('output_part2.wav')
files.download('output_part3.wav')
files.download('output_part4.wav')


# Gerekli kütüphaneleri yükleyelim
!pip install git+https://github.com/openai/whisper
!pip install python-docx

!pip install -q -U google-generativeai  

import whisper
import google.generativeai as genai
from docx import Document

# Whisper modelini yükle 
model = whisper.load_model("large-v3-turbo")

# İşlenecek 4 ses dosyası
audio_files = ["output_part1.wav", "output_part2.wav", "output_part3.wav", "output_part4.wav"]

# Tüm metinleri tutacak liste
all_texts = []

# Her ses dosyası için transkript ve tekil .docx oluşturma
for file in audio_files:
    print(f"Transkripte başlandı: {file}")
    result = model.transcribe(file, language="en", fp16=False)
    text = result["text"]
    all_texts.append(text)

    # Her biri için ayrı Word dosyası
    doc = Document()
    doc.add_heading(f"{file} - Transcript", level=1)
    doc.add_paragraph(text)
    single_docx = file.replace(".wav", "_transcript.docx")
    doc.save(single_docx)
    print(f"{single_docx} kaydedildi ✅")

# Tüm metinleri birleştiren tek bir Word dosyası
combined_doc = Document()
combined_doc.add_heading("Combined Transcript of All Audio Parts", level=1)

for idx, text in enumerate(all_texts):
    combined_doc.add_heading(f"Part {idx+1}", level=2)
    combined_doc.add_paragraph(text)

combined_doc.save("combined_transcript.docx")
print("\n✅ combined_transcript.docx kaydedildi.")

